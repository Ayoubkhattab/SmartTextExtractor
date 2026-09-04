"""MainWindow tests use a fake OcrWorkerPool/ScannerService (duck-typed,
same pattern as tests/scanner/test_service.py and
tests/concurrency/test_ocr_worker_pool.py) — real OCR is exercised by
tests/ocr/, real threading by tests/concurrency/; this suite is only
responsible for proving the Qt wiring itself is correct.
"""
from __future__ import annotations

from pathlib import Path

import pytest

from PyQt6.QtCore import Qt
from PyQt6.QtGui import QFont, QTextCursor

from smart_text_extractor.core.models import Document, DocumentUnit, OcrResult, OcrStatus, SourceType, TextSegment
from smart_text_extractor.scanner.models import ScannerDeviceInfo
from smart_text_extractor.ui.main_window import (
    _HEADING_FONT_POINT_SIZE,
    _LOW_CONFIDENCE_COLOR,
    _VERY_LOW_CONFIDENCE_COLOR,
    MainWindow,
)


class _FakeOcrPool:
    """submit() calls on_done synchronously — no threads, deterministic."""

    def __init__(self, result: OcrResult | None = None, error: Exception | None = None) -> None:
        self.result = result or OcrResult(raw_text="fake extracted text")
        self.error = error
        self.submitted_pages = []

    def submit(self, page, on_done):
        self.submitted_pages.append(page)
        if self.error is not None:
            page.ocr_status = OcrStatus.FAILED  # mirrors the real OcrWorkerPool's error path
            on_done(page, self.error)
        else:
            page.ocr_result = self.result
            page.ocr_status = OcrStatus.DONE
            on_done(page, None)


class _FakeScannerService:
    def __init__(self, devices: list[ScannerDeviceInfo] | None = None) -> None:
        self._devices = devices or []

    def discover(self):
        return self._devices


@pytest.fixture()
def document(tmp_path: Path) -> Document:
    return Document(source_type=SourceType.UPLOAD_IMAGE, temp_dir_path=tmp_path)


@pytest.fixture()
def sample_image(tmp_path: Path) -> Path:
    from PIL import Image

    path = tmp_path / "sample.png"
    Image.new("RGB", (100, 50), "white").save(path)
    return path


def test_window_starts_with_no_page_selected(qtbot, document) -> None:
    window = MainWindow(document, _FakeOcrPool(), _FakeScannerService())
    qtbot.addWidget(window)

    assert window._page_list.count() == 0
    assert "لا توجد صفحة محددة" in window._image_label.text()


def test_add_page_submits_to_pool_and_marks_done(qtbot, document, sample_image) -> None:
    pool = _FakeOcrPool(result=OcrResult(raw_text="hello world"))
    window = MainWindow(document, pool, _FakeScannerService())
    qtbot.addWidget(window)

    window._add_page(sample_image)

    assert len(document.pages) == 1
    assert document.pages[0] in pool.submitted_pages
    assert document.pages[0].ocr_status is OcrStatus.DONE
    assert window._page_list.count() == 1
    assert "تم" in window._page_list.item(0).text()


def test_add_page_shows_failure_in_list(qtbot, document, sample_image) -> None:
    pool = _FakeOcrPool(error=RuntimeError("boom"))
    window = MainWindow(document, pool, _FakeScannerService())
    qtbot.addWidget(window)

    window._add_page(sample_image)

    assert document.pages[0].ocr_status is OcrStatus.FAILED
    assert "فشل" in window._page_list.item(0).text()


def test_selecting_a_done_page_shows_its_extracted_text(qtbot, document, sample_image) -> None:
    pool = _FakeOcrPool(result=OcrResult(raw_text="the extracted text"))
    window = MainWindow(document, pool, _FakeScannerService())
    qtbot.addWidget(window)

    window._add_page(sample_image)
    window._page_list.setCurrentRow(0)

    assert window._text_edit.toPlainText() == "the extracted text"
    assert window._text_edit.isReadOnly() is False


def _char_format_at(text_edit, position: int):
    cursor = QTextCursor(text_edit.document())
    cursor.setPosition(position)
    cursor.setPosition(position + 1, QTextCursor.MoveMode.KeepAnchor)
    return cursor.charFormat()


def test_low_and_very_low_confidence_words_are_highlighted_distinctly(qtbot, document, sample_image) -> None:
    segments = [
        TextSegment("Good", 95.0),  # >= 75: no highlight
        TextSegment(" ", None),  # separator: never highlighted regardless of neighbors
        TextSegment("Iffy", 60.0),  # 50-75: low-confidence highlight
        TextSegment(" ", None),
        TextSegment("Bad", 40.0),  # < 50: very-low-confidence highlight
    ]
    result = OcrResult(raw_text="Good Iffy Bad", segments=segments)
    pool = _FakeOcrPool(result=result)
    window = MainWindow(document, pool, _FakeScannerService())
    qtbot.addWidget(window)

    window._add_page(sample_image)
    window._page_list.setCurrentRow(0)

    assert window._text_edit.toPlainText() == "Good Iffy Bad"
    assert _char_format_at(window._text_edit, 0).background().color().name() != _LOW_CONFIDENCE_COLOR.name()
    assert _char_format_at(window._text_edit, 5).background().color().name() == _LOW_CONFIDENCE_COLOR.name()
    assert _char_format_at(window._text_edit, 10).background().color().name() == _VERY_LOW_CONFIDENCE_COLOR.name()


def _cell_text(table, row: int, col: int) -> str:
    cell = table.cellAt(row, col)
    cursor = cell.firstCursorPosition()
    cursor.setPosition(cell.lastCursorPosition().position(), QTextCursor.MoveMode.KeepAnchor)
    return cursor.selectedText()


def _first_table(text_edit):
    # cursor.currentTable() on a cursor moved to document Start lands at
    # the table's own frame boundary, not "inside" a cell, and returns
    # None — walking the document's frame tree instead reliably finds
    # the table object itself (confirmed: PyQt returns the real QTextTable
    # subtype from childFrames(), not just a base QTextFrame).
    return text_edit.document().rootFrame().childFrames()[0]


def _open_page_with_document_units(qtbot, document, sample_image, units):
    result = OcrResult(raw_text="placeholder", document_units=units)
    pool = _FakeOcrPool(result=result)
    window = MainWindow(document, pool, _FakeScannerService())
    qtbot.addWidget(window)
    window._add_page(sample_image)
    window._page_list.setCurrentRow(0)
    return window


def test_table_document_unit_renders_as_a_real_table_grid(qtbot, document, sample_image) -> None:
    """A real user request (docs/phases/phase-2-ocr-pipeline.md): the
    live extracted-text panel should look organized like the original
    page, not a flat text dump — a table unit becomes an actual
    QTextTable grid, not pipe-separated text."""
    rows = [
        [[TextSegment("A", 90.0)], [TextSegment("B", 90.0)]],
        [[TextSegment("1", 90.0)], [TextSegment("2", 90.0)]],
    ]
    units = [DocumentUnit(kind="table", rows=rows)]
    window = _open_page_with_document_units(qtbot, document, sample_image, units)

    table = _first_table(window._text_edit)

    assert table is not None
    assert table.rows() == 2
    assert table.columns() == 2


def test_table_cells_render_in_right_to_left_column_order(qtbot, document, sample_image) -> None:
    """The real, empirically-confirmed regression this guards against
    (docs/phases/phase-2-ocr-pipeline.md): QTextTableFormat's
    layoutDirection does NOT reorder columns on its own — a table built
    without the explicit reversal put the first-to-read cell on the
    LEFT, which is backwards for Arabic. rows[0] is the first cell to
    read, so it must land in the RIGHTMOST visual column (columns - 1),
    not visual column 0.
    """
    rows = [[[TextSegment("first", 90.0)], [TextSegment("second", 90.0)], [TextSegment("third", 90.0)]]]
    units = [DocumentUnit(kind="table", rows=rows)]
    window = _open_page_with_document_units(qtbot, document, sample_image, units)

    table = _first_table(window._text_edit)

    assert _cell_text(table, 0, 2) == "first"  # rightmost visual column = first to read
    assert _cell_text(table, 0, 1) == "second"
    assert _cell_text(table, 0, 0) == "third"  # leftmost visual column = last to read


def test_ragged_table_row_leaves_the_correct_visual_cell_empty(qtbot, document, sample_image) -> None:
    """A short row (real messy tables have these — see _block_is_tabular)
    is padded at the END of reading order, which for RTL is the LEFTMOST
    visual column — verified explicitly rather than assumed, since this
    is exactly the kind of off-by-one a manual index reversal can get
    wrong silently."""
    rows = [
        [[TextSegment("A", 90.0)], [TextSegment("B", 90.0)], [TextSegment("C", 90.0)]],
        [[TextSegment("1", 90.0)]],  # short row: only 1 of 3 columns
    ]
    units = [DocumentUnit(kind="table", rows=rows)]
    window = _open_page_with_document_units(qtbot, document, sample_image, units)

    table = _first_table(window._text_edit)

    assert _cell_text(table, 1, 2) == "1"  # the only real cell, rightmost (first-to-read position)
    assert _cell_text(table, 1, 1) == ""
    assert _cell_text(table, 1, 0) == ""


def test_heading_document_unit_renders_bold_and_larger(qtbot, document, sample_image) -> None:
    units = [DocumentUnit(kind="heading", segments=[TextSegment("عنوان القسم", 90.0)])]
    window = _open_page_with_document_units(qtbot, document, sample_image, units)

    char_format = _char_format_at(window._text_edit, 0)
    assert char_format.fontWeight() == QFont.Weight.Bold
    assert char_format.fontPointSize() == _HEADING_FONT_POINT_SIZE


def test_low_confidence_word_inside_a_heading_still_gets_highlighted(qtbot, document, sample_image) -> None:
    """Confidence highlighting and heading styling combine — a heading
    is not automatically trustworthy just because it was recognized as
    one (see _char_format_for's docstring)."""
    units = [DocumentUnit(kind="heading", segments=[TextSegment("مشكوك", 40.0)])]
    window = _open_page_with_document_units(qtbot, document, sample_image, units)

    char_format = _char_format_at(window._text_edit, 0)
    assert char_format.fontWeight() == QFont.Weight.Bold  # still a heading
    assert char_format.background().color().name() == _VERY_LOW_CONFIDENCE_COLOR.name()  # and still flagged


def test_document_units_take_priority_over_flat_segments(qtbot, document, sample_image) -> None:
    """document_units, when present, is used instead of the flatter
    segments fallback — verified by rendering a heading (which only the
    document_units path produces) and checking a real Qt table object is
    NOT what's used for the plain-segments case (see
    test_low_and_very_low_confidence_words_are_highlighted_distinctly
    for the segments-only path this must not break)."""
    result = OcrResult(
        raw_text="placeholder",
        segments=[TextSegment("fallback text, should not be shown", None)],
        document_units=[DocumentUnit(kind="heading", segments=[TextSegment("العنوان الحقيقي", 90.0)])],
    )
    pool = _FakeOcrPool(result=result)
    window = MainWindow(document, pool, _FakeScannerService())
    qtbot.addWidget(window)

    window._add_page(sample_image)
    window._page_list.setCurrentRow(0)

    assert window._text_edit.toPlainText() == "العنوان الحقيقي"


def test_editing_text_calls_set_edited_text_not_raw_text(qtbot, document, sample_image) -> None:
    pool = _FakeOcrPool(result=OcrResult(raw_text="original"))
    window = MainWindow(document, pool, _FakeScannerService())
    qtbot.addWidget(window)

    window._add_page(sample_image)
    window._page_list.setCurrentRow(0)

    window._text_edit.setPlainText("edited by the user")

    page = document.pages[0]
    assert page.ocr_result.raw_text == "original"
    assert page.ocr_result.edited_text == "edited by the user"


def test_scan_with_no_devices_shows_information_message(qtbot, document, monkeypatch) -> None:
    window = MainWindow(document, _FakeOcrPool(), _FakeScannerService(devices=[]))
    qtbot.addWidget(window)

    shown = []
    monkeypatch.setattr(
        "smart_text_extractor.ui.main_window.QMessageBox.information",
        lambda *args, **kwargs: shown.append(args),
    )

    window._on_scan()

    assert len(shown) == 1


def test_scan_with_devices_shows_a_different_message(qtbot, document, monkeypatch) -> None:
    devices = [ScannerDeviceInfo(device_id="dev-1", name="Fake Scanner")]
    window = MainWindow(document, _FakeOcrPool(), _FakeScannerService(devices=devices))
    qtbot.addWidget(window)

    shown = []
    monkeypatch.setattr(
        "smart_text_extractor.ui.main_window.QMessageBox.information",
        lambda *args, **kwargs: shown.append(args),
    )

    window._on_scan()

    assert len(shown) == 1
    assert "1" in shown[0][2]  # message text mentions the device count


def test_export_markdown_with_no_done_pages_shows_information_message(qtbot, document, monkeypatch) -> None:
    window = MainWindow(document, _FakeOcrPool(), _FakeScannerService())
    qtbot.addWidget(window)

    shown = []
    monkeypatch.setattr(
        "smart_text_extractor.ui.main_window.QMessageBox.information",
        lambda *args, **kwargs: shown.append(args),
    )
    save_calls = []
    monkeypatch.setattr(
        "smart_text_extractor.ui.main_window.QFileDialog.getSaveFileName",
        lambda *args, **kwargs: save_calls.append(args) or ("", ""),
    )

    window._on_export_markdown()

    assert len(shown) == 1
    assert save_calls == []  # never even prompted for a save path


def test_export_markdown_writes_markdown_for_unedited_pages_and_plain_text_for_edited_ones(
    qtbot, document, sample_image, tmp_path: Path, monkeypatch
) -> None:
    pool = _FakeOcrPool()  # submit() assigns the SAME OcrResult instance to every page
    window = MainWindow(document, pool, _FakeScannerService())
    qtbot.addWidget(window)

    window._add_page(sample_image)  # page 0: kept as a fresh, unedited OCR result below
    window._add_page(sample_image)  # page 1: will be edited by the user below

    # Give each page its own independent OcrResult — the fake pool shares
    # one instance across every submit() call, which isn't representative
    # of the real per-page OcrWorkerPool and would make the two pages'
    # edited_text alias each other if left as-is.
    document.pages[0].ocr_result = OcrResult(raw_text="RAW A", markdown="## RAW A (as markdown)")
    document.pages[1].ocr_result = OcrResult(raw_text="RAW B", markdown="## RAW B (as markdown)")

    window._page_list.setCurrentRow(1)
    window._text_edit.setPlainText("user's edited text")  # sets page 1's edited_text

    save_path = tmp_path / "export"  # deliberately no .md suffix
    monkeypatch.setattr(
        "smart_text_extractor.ui.main_window.QFileDialog.getSaveFileName",
        lambda *args, **kwargs: (str(save_path), ""),
    )

    window._on_export_markdown()

    written_path = tmp_path / "export.md"
    assert written_path.exists()
    content = written_path.read_text(encoding="utf-8")
    assert content == "## RAW A (as markdown)\n\n---\n\nuser's edited text"


def test_export_markdown_skips_pages_excluded_from_range(qtbot, document, sample_image, tmp_path: Path, monkeypatch) -> None:
    pool = _FakeOcrPool(result=OcrResult(raw_text="included", markdown="included"))
    window = MainWindow(document, pool, _FakeScannerService())
    qtbot.addWidget(window)

    window._add_page(sample_image)
    document.pages[0].included_in_range = False

    save_path = tmp_path / "export.md"
    monkeypatch.setattr(
        "smart_text_extractor.ui.main_window.QFileDialog.getSaveFileName",
        lambda *args, **kwargs: (str(save_path), ""),
    )
    shown = []
    monkeypatch.setattr(
        "smart_text_extractor.ui.main_window.QMessageBox.information",
        lambda *args, **kwargs: shown.append(args),
    )

    window._on_export_markdown()

    assert len(shown) == 1  # no exportable pages left once the only page is excluded
    assert not save_path.exists()


def test_export_word_writes_a_real_docx_using_document_units(qtbot, document, sample_image, tmp_path: Path, monkeypatch) -> None:
    import docx

    units = [
        DocumentUnit(kind="heading", segments=[TextSegment("عنوان", 90.0)]),
        DocumentUnit(kind="paragraph", segments=[TextSegment("نص الفقرة", 90.0)]),
    ]
    pool = _FakeOcrPool(result=OcrResult(raw_text="RAW", document_units=units))
    window = MainWindow(document, pool, _FakeScannerService())
    qtbot.addWidget(window)

    window._add_page(sample_image)

    save_path = tmp_path / "export"  # no .docx suffix — MainWindow must add one
    monkeypatch.setattr(
        "smart_text_extractor.ui.main_window.QFileDialog.getSaveFileName",
        lambda *args, **kwargs: (str(save_path), ""),
    )

    window._on_export_word()

    written_path = tmp_path / "export.docx"
    assert written_path.exists()
    reopened = docx.Document(str(written_path))
    assert any(p.text == "عنوان" and p.style.name == "Heading 2" for p in reopened.paragraphs)
    assert any(p.text == "نص الفقرة" for p in reopened.paragraphs)


def test_opening_a_multi_page_pdf_adds_one_page_per_pdf_page(qtbot, document, tmp_path: Path) -> None:
    import pymupdf

    pdf_path = tmp_path / "doc.pdf"
    doc = pymupdf.open()
    for i in range(3):
        page = doc.new_page(width=300, height=200)
        page.insert_text((20, 20), f"page {i + 1}", fontsize=16)
    doc.save(str(pdf_path))
    doc.close()

    pool = _FakeOcrPool()
    window = MainWindow(document, pool, _FakeScannerService())
    qtbot.addWidget(window)

    window._open_file(pdf_path)

    assert len(document.pages) == 3
    assert window._page_list.count() == 3
    assert len(pool.submitted_pages) == 3


def test_add_page_with_a_native_result_skips_ocr_entirely(qtbot, document, sample_image) -> None:
    """_add_page's native_result path (§7.1 extension) is DONE
    immediately from that text, never reaching the OCR pool — the
    underlying mechanism, kept and tested even though _open_file no
    longer wires it in automatically for PDFs.

    REAL-WORLD REGRESSION (docs/phases/phase-2-ocr-pipeline.md): this
    used to be wired into _open_file via import_pdf_pages for every PDF.
    Reverted after live testing on a real user document: PyMuPDF's own
    text extraction (page.get_text) came back character-corrupted for
    specific letter sequences on that PDF (e.g. "الافتراضية" -> "الفرتاضية")
    — confirmed via a direct OCR comparison on the same real page that
    Tesseract read those exact words correctly where the embedded text
    layer did not, i.e. the PDF's own font encoding is broken for that
    document, not something our extraction code causes or can detect
    reliably yet. OCR is the safe default until there's a validated way
    to tell a trustworthy embedded text layer from a corrupted one.
    """
    native_result = OcrResult(raw_text="word0 word1", confidence_score=100.0)
    pool = _FakeOcrPool()
    window = MainWindow(document, pool, _FakeScannerService())
    qtbot.addWidget(window)

    window._add_page(sample_image, native_result)

    assert len(document.pages) == 1
    assert pool.submitted_pages == []  # never went through OCR
    assert document.pages[0].ocr_status is OcrStatus.DONE
    assert document.pages[0].ocr_result.confidence_score == 100.0
    assert "word0" in document.pages[0].ocr_result.raw_text
    assert "نص أصلي" in window._page_list.item(0).text()


def test_opening_a_pdf_always_goes_through_ocr_even_with_a_native_text_layer(
    qtbot, document, tmp_path: Path
) -> None:
    """_open_file does not use import_pdf_pages/native extraction — see
    test_add_page_with_a_native_result_skips_ocr_entirely's docstring for
    why this was reverted. Every PDF page, regardless of its own embedded
    text, goes through the OCR pool exactly as an image-only page would.
    """
    import pymupdf

    pdf_path = tmp_path / "doc.pdf"
    doc = pymupdf.open()
    page = doc.new_page(width=600, height=400)
    text = " ".join(f"word{i}" for i in range(20))  # well over the native-text word threshold
    page.insert_textbox(pymupdf.Rect(20, 20, 580, 380), text, fontsize=14)
    doc.save(str(pdf_path))
    doc.close()

    pool = _FakeOcrPool()
    window = MainWindow(document, pool, _FakeScannerService())
    qtbot.addWidget(window)

    window._open_file(pdf_path)

    assert len(document.pages) == 1
    assert len(pool.submitted_pages) == 1  # went through OCR despite having a real text layer


def test_opening_a_broken_pdf_shows_a_warning_not_a_crash(qtbot, document, tmp_path: Path, monkeypatch) -> None:
    fake_pdf = tmp_path / "not_really_a_pdf.pdf"
    fake_pdf.write_bytes(b"this is not a valid pdf file")

    window = MainWindow(document, _FakeOcrPool(), _FakeScannerService())
    qtbot.addWidget(window)

    shown = []
    monkeypatch.setattr(
        "smart_text_extractor.ui.main_window.QMessageBox.warning",
        lambda *args, **kwargs: shown.append(args),
    )

    window._open_file(fake_pdf)

    assert len(shown) == 1
    assert len(document.pages) == 0


class TestPageManagementWiring:
    """US-07 (reorder + undo), US-08 (page range) and manual retry (§3.2)
    all had working, tested domain logic with no way for a user to reach
    it. These prove the UI actually drives that logic."""

    def _window_with_pages(self, qtbot, document, tmp_path, count=3):
        from PIL import Image

        pool = _FakeOcrPool()
        window = MainWindow(document, pool, _FakeScannerService())
        qtbot.addWidget(window)
        for i in range(count):
            path = tmp_path / f"page{i}.png"
            Image.new("RGB", (60, 40), "white").save(path)
            window._add_page(path)
        return window

    def test_dragging_a_page_reorders_the_document(self, qtbot, document, tmp_path) -> None:
        window = self._window_with_pages(qtbot, document, tmp_path)
        original = [page.id for page in document.pages]

        # what Qt leaves behind after an InternalMove drag: the list rows
        # are already in the new order, then rowsMoved fires.
        item = window._page_list.takeItem(0)
        window._page_list.insertItem(2, item)
        window._on_rows_moved()

        assert [page.id for page in document.pages] == [original[1], original[2], original[0]]

    def test_reordering_a_page_that_is_mid_ocr_is_refused_and_the_list_is_restored(
        self, qtbot, document, tmp_path
    ) -> None:
        window = self._window_with_pages(qtbot, document, tmp_path)
        original = [page.id for page in document.pages]
        document.pages[0].ocr_status = OcrStatus.PROCESSING  # §3.1: locked

        item = window._page_list.takeItem(0)
        window._page_list.insertItem(2, item)
        window._on_rows_moved()

        assert [page.id for page in document.pages] == original  # document untouched
        listed = [window._page_list.item(r).data(Qt.ItemDataRole.UserRole) for r in range(3)]
        assert listed == original  # and the list was put back to match it

    def test_undo_restores_the_previous_order(self, qtbot, document, tmp_path) -> None:
        window = self._window_with_pages(qtbot, document, tmp_path)
        original = [page.id for page in document.pages]

        item = window._page_list.takeItem(0)
        window._page_list.insertItem(2, item)
        window._on_rows_moved()
        window._on_undo_reorder()

        assert [page.id for page in document.pages] == original
        assert [window._page_list.item(r).data(Qt.ItemDataRole.UserRole) for r in range(3)] == original

    def test_unchecking_a_page_excludes_it_from_export(self, qtbot, document, tmp_path) -> None:
        window = self._window_with_pages(qtbot, document, tmp_path)
        assert len(window._exportable_pages()) == 3

        window._page_list.item(1).setCheckState(Qt.CheckState.Unchecked)

        assert document.pages[1].included_in_range is False
        assert len(window._exportable_pages()) == 2

    def test_retry_is_offered_only_for_a_failed_page_and_resubmits_it(self, qtbot, document, tmp_path) -> None:
        from PIL import Image

        pool = _FakeOcrPool(error=RuntimeError("boom"))
        window = MainWindow(document, pool, _FakeScannerService())
        qtbot.addWidget(window)
        path = tmp_path / "failing.png"
        Image.new("RGB", (60, 40), "white").save(path)
        window._add_page(path)

        assert window._retry_action.isEnabled() is True
        submitted_before = len(pool.submitted_pages)

        window._on_retry_page()

        assert len(pool.submitted_pages) == submitted_before + 1
