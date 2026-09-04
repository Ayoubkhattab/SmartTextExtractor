"""Word (.docx) export (§7.3, US-09) via python-docx.

Consumes the same DocumentUnit structure the Markdown export uses
(smart_text_extractor.ocr.reorder.classify_document_units, via
OcrResult.document_units) so heading/table classification (§7.1.1) lives
in exactly one place, not reimplemented per export format.
"""
from __future__ import annotations

from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from smart_text_extractor.core.models import DocumentUnit, PageLayout, TextSegment

# A page's content is either its still-structured OCR result (most pages)
# or a plain string — a page the user has edited only has edited_text, a
# plain string with no positional data left to classify into headings/
# tables (see MainWindow._on_export_markdown's identical distinction).
PageContent = list[DocumentUnit] | str


def _set_rtl(paragraph: Paragraph) -> None:
    """Marks a paragraph right-to-left — both visually (right alignment)
    and structurally (the real w:bidi OOXML flag, verified via a
    round-trip: python-docx reads it back correctly after save/reopen).
    Plain right-alignment alone is a cosmetic-only fix; w:bidi is what
    actually tells Word this paragraph's text direction and cursor
    behavior are RTL, which matters for a tool whose primary content is
    Arabic.
    """
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph_properties = paragraph._p.get_or_add_pPr()
    paragraph_properties.append(paragraph_properties.makeelement(qn("w:bidi"), {}))


def _set_table_rtl(table: docx.table.Table) -> None:
    """Marks a table right-to-left via the real w:bidiVisual OOXML flag —
    verified visually (rendered through actual Word to PDF): without
    this, python-docx's table.cell(0, 0) always lands in the leftmost
    visual column regardless of paragraph direction, so a table built
    from already-correct right-to-left cell content (row[0] = the first
    cell to read) rendered with its columns in the wrong order — cell 0
    on the left instead of the right. w:bidiVisual is what actually
    flips the visual column order to match; per-paragraph w:bidi/right
    alignment inside each cell only affects that cell's own text, not
    which physical column it appears in.
    """
    table_properties = table._tbl.tblPr
    table_properties.append(table_properties.makeelement(qn("w:bidiVisual"), {}))


def _segments_to_text(segments: list[TextSegment]) -> str:
    return "".join(segment.text for segment in segments)


def _add_styled_runs(paragraph: Paragraph, segments: list[TextSegment]) -> None:
    """Writes the segments as Word runs that keep the source page's own
    look — real font size, weight and colour — instead of flattening
    everything to the default body style.

    Only a PDF text layer records any of that (models.TextStyle); an OCR
    page's segments carry none and come out as plain runs exactly as
    before. Consecutive segments sharing a style are written as ONE run,
    so a normal paragraph stays a single run rather than one per word.
    """
    if not segments:
        return

    current_style = segments[0].style
    buffer: list[str] = []

    def flush() -> None:
        if not buffer:
            return
        run = paragraph.add_run("".join(buffer))
        if current_style is not None:
            if current_style.font_size:
                run.font.size = Pt(current_style.font_size)
            if current_style.bold:
                run.bold = True
            if current_style.italic:
                run.italic = True
            if current_style.color:
                run.font.color.rgb = RGBColor.from_string(current_style.color.lstrip("#").upper())

    for segment in segments:
        if segment.style != current_style:
            flush()
            buffer = []
            current_style = segment.style
        buffer.append(segment.text)
    flush()


def _add_table(document: docx.document.Document, rows: list[list[list[TextSegment]]]) -> None:
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    _set_table_rtl(table)
    for row_index, row in enumerate(rows):
        for column_index in range(column_count):
            cell = table.cell(row_index, column_index)
            cell.text = _segments_to_text(row[column_index]) if column_index < len(row) else ""
            for paragraph in cell.paragraphs:
                _set_rtl(paragraph)


def _add_units(document: docx.document.Document, units: list[DocumentUnit]) -> None:
    for unit in units:
        if unit.kind == "table":
            _add_table(document, unit.rows)
            continue

        # add_heading()/add_paragraph() with no text, then styled runs: the
        # heading STYLE still marks it as a heading for Word's navigation
        # and table of contents, while the runs carry the source page's own
        # size/weight/colour on top of it.
        paragraph = document.add_heading("", level=2) if unit.kind == "heading" else document.add_paragraph()
        _set_rtl(paragraph)
        if unit.alignment == "center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if unit.space_before_points:
            # Reproduces the blank space that sat above this block on the
            # source page, instead of Word's uniform paragraph gap.
            paragraph.paragraph_format.space_before = Pt(unit.space_before_points)
        _add_styled_runs(paragraph, unit.segments)


def _add_plain_text(document: docx.document.Document, text: str) -> None:
    """Fallback for a page the user has edited (see PageContent) — one
    Word paragraph per line, since a single Word paragraph does not treat
    an embedded "\\n" as a line break the way plain text does."""
    for line in text.split("\n"):
        _set_rtl(document.add_paragraph(line))


def _apply_page_layout(document: docx.document.Document, layout: PageLayout) -> None:
    """Puts the Word document on the same paper, with the same margins, as
    the page it came from.

    This is the single biggest thing standing between an export and "no
    re-formatting needed". Word's default is US Letter with 1in margins; a
    document that was A4 with a 70pt left margin re-flows every single line
    against that, so line breaks, page breaks and anything positioned
    relative to them all move. Measured on this project's own files, one
    source is A4 and another is Letter, so the size is taken from the
    source rather than defaulted either way.
    """
    for section in document.sections:
        section.page_width = Pt(layout.width_points)
        section.page_height = Pt(layout.height_points)
        section.left_margin = Pt(layout.margin_left)
        section.right_margin = Pt(layout.margin_right)
        section.top_margin = Pt(layout.margin_top)
        section.bottom_margin = Pt(layout.margin_bottom)


def build_docx(pages: list[PageContent], page_layout: PageLayout | None = None) -> docx.document.Document:
    """pages: one entry per exported page, in export order. A page
    boundary becomes a real Word page break, not just blank space — Word
    has that primitive and Markdown doesn't.

    page_layout, when given, is the source page geometry to reproduce (see
    _apply_page_layout). It applies to the whole document because a Word
    section spans pages; the documents this exports are single-geometry, so
    the first page's layout is the document's layout.
    """
    document = docx.Document()
    if page_layout is not None:
        _apply_page_layout(document, page_layout)
    for page_index, content in enumerate(pages):
        if page_index > 0:
            document.add_page_break()
        if isinstance(content, str):
            _add_plain_text(document, content)
        else:
            _add_units(document, content)
    return document


def export_docx(pages: list[PageContent], path: Path, page_layout: PageLayout | None = None) -> None:
    build_docx(pages, page_layout).save(str(path))
