from __future__ import annotations

from pathlib import Path

import docx

from smart_text_extractor.core.models import DocumentUnit, TextSegment
from smart_text_extractor.export.docx_export import build_docx, export_docx


def _has_page_break(paragraph) -> bool:
    return 'type="page"' in paragraph._p.xml


def _seg(text: str, confidence: float = 90.0) -> list[TextSegment]:
    return [TextSegment(text, confidence)]


def test_paragraph_unit_becomes_an_rtl_right_aligned_paragraph() -> None:
    document = build_docx([[DocumentUnit(kind="paragraph", segments=_seg("مرحباً بكم"))]])

    paragraphs = [p for p in document.paragraphs if p.text]
    assert len(paragraphs) == 1
    assert paragraphs[0].text == "مرحباً بكم"
    assert 'w:bidi' in paragraphs[0]._p.xml


def test_heading_unit_becomes_a_heading_2_style_paragraph() -> None:
    document = build_docx([[DocumentUnit(kind="heading", segments=_seg("عنوان القسم"))]])

    headings = [p for p in document.paragraphs if p.style.name == "Heading 2"]
    assert len(headings) == 1
    assert headings[0].text == "عنوان القسم"


def test_table_unit_becomes_a_real_docx_table_with_matching_cells() -> None:
    rows = [[_seg("A"), _seg("B"), _seg("C")], [_seg("1"), _seg("2")]]  # ragged, like a real messy table
    unit = DocumentUnit(kind="table", rows=rows)

    document = build_docx([[unit]])

    assert len(document.tables) == 1
    table = document.tables[0]
    assert len(table.rows) == 2
    assert len(table.columns) == 3  # padded to the widest row
    # w:bidiVisual (real-world regression, docs/phases/phase-2-ocr-pipeline.md):
    # without it, a real Word render put cell 0 in the leftmost visual
    # column regardless of paragraph direction — verified by rendering
    # through actual Word to PDF, with and without this flag.
    assert "bidiVisual" in table._tbl.tblPr.xml
    assert [cell.text for cell in table.rows[0].cells] == ["A", "B", "C"]
    assert [cell.text for cell in table.rows[1].cells] == ["1", "2", ""]  # short row padded with an empty cell


def test_multiple_pages_get_a_real_page_break_between_them() -> None:
    document = build_docx(
        [
            [DocumentUnit(kind="paragraph", segments=_seg("page one"))],
            [DocumentUnit(kind="paragraph", segments=_seg("page two"))],
        ]
    )

    assert any(_has_page_break(p) for p in document.paragraphs)


def test_single_page_has_no_page_break() -> None:
    document = build_docx([[DocumentUnit(kind="paragraph", segments=_seg("only page"))]])

    assert not any(_has_page_break(p) for p in document.paragraphs)


def test_edited_page_falls_back_to_one_paragraph_per_line() -> None:
    """A page the user has edited only has plain edited_text left (no
    positional data for headings/tables — see MainWindow._on_export_word),
    so PageContent accepts a plain string for that page instead of
    list[DocumentUnit]."""
    document = build_docx(["line one\nline two"])

    texts = [p.text for p in document.paragraphs]
    assert texts == ["line one", "line two"]


def test_export_docx_writes_a_file_that_reopens_with_the_same_content(tmp_path: Path) -> None:
    path = tmp_path / "out.docx"

    export_docx([[DocumentUnit(kind="paragraph", segments=_seg("نص للتحقق"))]], path)

    assert path.exists()
    reopened = docx.Document(str(path))
    assert any(p.text == "نص للتحقق" for p in reopened.paragraphs)


class TestSourcePageStylingInWord:
    """The Word export should look like the page it came from, not like the
    default body style — real font size, weight, colour and centring."""

    def _paragraph(self, units):
        return build_docx([units]).paragraphs

    def test_font_size_and_weight_reach_the_word_run(self) -> None:
        from docx.shared import Pt

        from smart_text_extractor.core.models import TextStyle

        units = [
            DocumentUnit(
                kind="heading",
                segments=[TextSegment("عنوان", 100.0, TextStyle(font_size=24.0, bold=True))],
            )
        ]

        run = self._paragraph(units)[0].runs[0]

        assert run.font.size == Pt(24.0)
        assert run.bold is True

    def test_text_colour_reaches_the_word_run(self) -> None:
        from docx.shared import RGBColor

        from smart_text_extractor.core.models import TextStyle

        units = [DocumentUnit(kind="paragraph", segments=[TextSegment("نص", 100.0, TextStyle(color="#333333"))])]

        assert self._paragraph(units)[0].runs[0].font.color.rgb == RGBColor.from_string("333333")

    def test_a_centred_unit_is_centred_in_word(self) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        units = [DocumentUnit(kind="heading", segments=[TextSegment("عنوان", 100.0)], alignment="center")]

        assert self._paragraph(units)[0].alignment == WD_ALIGN_PARAGRAPH.CENTER

    def test_a_heading_keeps_its_word_heading_style(self) -> None:
        """The style is what puts it in Word's navigation pane; the run
        formatting sits on top of it rather than replacing it."""
        units = [DocumentUnit(kind="heading", segments=[TextSegment("عنوان", 100.0)])]

        assert self._paragraph(units)[0].style.name.startswith("Heading")

    def test_consecutive_segments_sharing_a_style_become_one_run(self) -> None:
        from smart_text_extractor.core.models import TextStyle

        style = TextStyle(font_size=14.0)
        units = [
            DocumentUnit(
                kind="paragraph",
                segments=[TextSegment("كلمة", 100.0, style), TextSegment(" ", None, style), TextSegment("أخرى", 100.0, style)],
            )
        ]

        runs = self._paragraph(units)[0].runs

        assert len(runs) == 1
        assert runs[0].text == "كلمة أخرى"

    def test_unstyled_ocr_segments_still_export(self) -> None:
        units = [DocumentUnit(kind="paragraph", segments=[TextSegment("plain", 90.0)])]

        assert self._paragraph(units)[0].text == "plain"
